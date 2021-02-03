# Arthor: B.X. Weinstein Xiao
# Contact: rchopin@outlook.com, bangxi_xiao@brown,edu

import requests
import numpy as np
import json
import os
from bs4 import BeautifulSoup
from random import choice
import docx
import pdfminer.high_level as ph
from win32com import client as wc


class SZSE(object):
    def __init__(self):
        self.ua_library = ['Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) ',
                           'Chrome/45.0.2454.85 Safari/537.36 115Browser/6.0.3',
                           'Mozilla/5.0 (Macintosh; U; Intel Mac OS X 10_6_8; en-us) AppleWebKit/534.50 (KHTML, like Gecko) Version/5.1 Safari/534.50',
                           'Mozilla/5.0 (Windows; U; Windows NT 6.1; en-us) AppleWebKit/534.50 (KHTML, like Gecko) Version/5.1 Safari/534.50',
                           'Mozilla/4.0 (compatible; MSIE 8.0; Windows NT 6.0; Trident/4.0)',
                           'Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 6.0)',
                           'Mozilla/5.0 (Windows NT 6.1; rv:2.0.1) Gecko/20100101 Firefox/4.0.1',
                           'Opera/9.80 (Windows NT 6.1; U; en) Presto/2.8.131 Version/11.11',
                           'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_7_0) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.56 Safari/535.11',
                           'Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; Trident/4.0; SE 2.X MetaSr 1.0; SE 2.X MetaSr 1.0; .NET CLR 2.0.50727; SE 2.X MetaSr 1.0)',
                           'Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; Trident/5.0',
                           'Mozilla/5.0 (Windows NT 6.1; rv:2.0.1) Gecko/20100101 Firefox/4.0.1',
                           'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.132 Safari/537.36']
        self.tab_indicator = {'主板': 1, '中小企业板': 2, '创业板': 3}
        self.rc = ['gsdm', 'gsjc', 'fhrq', 'hjlb', 'ck', 'hfck', 'unique_id']
        self.download_header = 'http://reportdocs.static.szse.cn/'
        self.cache_filename = "cache_requests_szse"

    def url_processor(self,
                      u: str):
        """
        :param u: lxml/html-formatted, url information included string
        :return: url
        """
        soup = BeautifulSoup(u, 'lxml')
        if soup.find_all('a'):
            r = self.download_header + soup.find_all('a')[0]['encode-open']
        else:
            r = ''
        return r

    def download_all(self,
                     url: str,
                     id: str,
                     file_path: str):
        if not url:
            return
        r = requests.get(url)

        fmt = url.split('/')[-1].split('.')[-1]
        with open(file_path + '\\{}.{}'.format(id, fmt), 'wb+') as f:
            f.write(r.content)

    def get_local_pdf_content(self,
                              file_name: str,
                              file_path: str):
        id = file_name.split('.')[0]
        try:
            t = ph.extract_text(file_path + '\\{}'.format(file_name))
        except:
            t = ''
        return [t, id]

    def get_local_docx_content(self,
                               file_name: str,
                               file_path: str):
        id = file_name.split('.')[0]
        t = ''.join(
            list(map(lambda x: x.text, docx.Document(file_path + '\\{}'.format(file_name)).paragraphs)))

        return [t, id]

    def get_local_doc_content(self,
                              file_name: str,
                              file_path: str):
        id = file_name.split('.')[0]
        word = wc.Dispatch('Word.Application')
        doc = word.Documents.Open(file_path + '\\{}'.format(file_name))
        doc.SaveAs(file_path + '\\{}'.format(file_name), 12, False, "", True, "", False, False,
                   False,
                   False)
        t = ''.join(list(
            map(lambda x: x.text, docx.Document(file_path + '\\{}'.format(file_name)).paragraphs)))
        return [t, id]

    # def get_pdf_content(self,
    #                     url: str):
    #     """
    #     :param url: target url to get pdf content
    #     :return: string-like text of content
    #     """
    #     if not url:
    #         return ''
    #
    #     r = requests.get(url)
    #     if 'pdf' in url:
    #         with open('./cache/{}.pdf'.format(self.cache_filename), 'wb+') as f:
    #             f.write(r.content)
    #         try:
    #             t = ph.extract_text('./cache/{}.pdf'.format(self.cache_filename).format(id))
    #         except:
    #             t = ''
    #     elif 'docx' in url:
    #         with open('./cache/{}.docx'.format(self.cache_filename), 'wb+') as f:
    #             f.write(r.content)
    #         try:
    #             t = ''.join(list(map(lambda x: x.text, docx.Document('./cache/{}.docx'.format(self.cache_filename)).paragraphs)))
    #         except:
    #             t = ''
    #     elif 'doc' in url:
    #         with open('./cache/{}.doc'.format(self.cache_filename), 'wb+') as f:
    #             f.write(r.content)
    #         try:
    #             word = wc.Dispatch('Word.Application')
    #             doc = word.Documents.Open('.\\cache\\{}.doc'.format(self.cache_filename))
    #             doc.SaveAs('.\\cache\\{}.docx'.format(self.cache_filename), 12, False, "", True, "", False, False,
    #                        False,
    #                        False)
    #             t = ''.join(list(
    #                 map(lambda x: x.text, docx.Document('.\\cache\\{}.docx'.format(self.cache_filename)).paragraphs)))
    #         except:
    #             t = ''
    #     else:
    #         print('Unabel to recognize file format from url {}'.format(url))
    #         t = ''
    #     return t

    def multi_decoder(self,
                      s: bytes):
        """
        :param s: bytes-format to be decoded with
        :return: decoded string
        """
        try:
            ds = str(s, 'utf_8_sig')
        except UnicodeDecodeError:
            try:
                ds = str(s, 'gbk')
            except UnicodeDecodeError:
                try:
                    ds = str(s, 'GB18030')
                except UnicodeDecodeError:
                    ds = str(s, 'gb2312', 'ignore')
        return ds

    def get_page(self,
                 p: int,
                 tab: str):
        """
        :param p: page number
        :param tab: type of block
        :return: required columns from scrapped data
        """
        if tab in self.tab_indicator.keys():
            tab = str(self.tab_indicator[tab])
        else:
            print('找不到对应板块，请在 主板 中小企业板 创业板 之间进行选择。')
        page_no = str(p)
        rand = str(np.random.randint(1, 100))
        url = '''http://www.szse.cn/api/report/ShowReport/data?SHOWTYPE=JSON&CATALOGID=main_wxhj&TABKEY=tab{}&PAGENO={}&random=0.6437020{}304172'''
        url = url.format(tab, page_no, rand)

        headers = {
            'User-Agent': choice(self.ua_library),
            'Referer': 'http://www.szse.cn/disclosure/supervision/inquire/index.html',
            'Host': 'www.szse.cn'}

        rsp = requests.get(url=url, headers=headers)
        rsp_str = self.multi_decoder(rsp.content)

        if len(rsp_str) < 500:
            print('Warning, scrapped file may be incorrect from page {}.'.format(page_no))
        else:
            pass

        json_str = json.loads(rsp_str)
        tgt_data = json_str[int(tab) - 1]['data']
        d = []
        for dd in tgt_data:
            dd['unique_id'] = str(np.random.random()).replace('.', '')
            dd['ck'] = self.url_processor(dd['ck'])
            dd['hfck'] = self.url_processor(dd['hfck'])
            d.append(dd)
        self.rc = d[0].keys()
        return d

    def extract_info(self,
                     d: list):
        """

        :param d: a list composed of dictionaries that need to extract information from
        :return: rc, extracted information
        """
        g_info = [[x.get(y) for y in self.rc] for x in d]

        return g_info